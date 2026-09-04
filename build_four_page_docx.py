"""Build the four-page editable manuscript in the supplied IEEE Word style."""

from __future__ import annotations

from copy import deepcopy
from io import BytesIO
from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "paper" / "main.docx"
OUTPUT = SOURCE
PROJECT_URL = "https://github.com/master-abc/WSR-YOLO"
ACCURACY_FIGURE = ROOT / "paper" / "figures" / "accuracy_evidence.png"
QUALITATIVE_FIGURE = ROOT / "paper" / "figures" / "qualitative_operating_point.png"
CONTEXT_FIGURE = ROOT / "paper" / "figures" / "context_route_evidence.png"
FALSE_ALARM_FIGURE = ROOT / "paper" / "figures" / "false_alarm_tradeoff.png"
ROBUSTNESS_FIGURE = ROOT / "paper" / "figures" / "robustness_frequency_compact.png"

TITLE_LINE_1 = "WSR-YOLO for PCB Defect Image Processing"
TITLE_LINE_2 = "with Wavelet-Conditioned Sparse Routing"
ACKNOWLEDGEMENT = (
    "The work was supported by the Teaching, Science, and Innovation Teaching "
    "and Learning Project of Guangdong University of Science and Technology "
    "under Grant (GKJXXZ2024008) and 2025 Undergraduate Innovation and "
    "Entrepreneurship Training Program Project under Grant (202513719002)."
)

ABSTRACT = (
    "Traditional image processing for industrial visual inspection relies on "
    "hand-crafted filtering, thresholding, and edge operators, whereas WSR-YOLO "
    "is a modern deep-learning approach. WSR ranks P3 features with fixed Haar cues "
    "and refines an exact top-k budget. On DsPCBSD+, seven paired runs give 46.69±0.72 "
    "AP50:95 versus 46.54±0.36 for YOLO11s (p=0.8125); three descriptive DeepPCB "
    "pairs gain 4.20 points but increase clean-template alarms. At a 25% budget, "
    "routes enrich box cells by 3.16×/3.55×, although spatial priors and matched "
    "convolution explain much of the gain. WSR adds 13.3K parameters and 0.038 "
    "GFLOPs but is 1.18–1.21× slower. Negative-aware training reduces DeepPCB "
    "board-FPR from 53.5% to 27.5%; an exploratory paired-reference detector "
    "reaches 1.4% and 0.959 F1 under a changed input contract."
)

REFERENCES = [
    "S. Ren et al., “Faster R-CNN: Towards Real-Time Object Detection with Region Proposal Networks,” NeurIPS, 2015. https://arxiv.org/abs/1506.01497",
    "N. Carion et al., “End-to-End Object Detection with Transformers,” ECCV, 2020. https://arxiv.org/abs/2005.12872",
    "P. Sun et al., “Sparse R-CNN: End-to-End Object Detection with Learnable Proposals,” CVPR, 2021. https://arxiv.org/abs/2011.12450",
    "G. Jocher and J. Qiu, “Ultralytics YOLO11,” version 11.0.0, 2024. https://github.com/ultralytics/ultralytics",
    "Y. Peng et al., “D-FINE: Redefine Regression Task in DETRs as Fine-grained Distribution Refinement,” ICLR, 2025. https://arxiv.org/abs/2410.13842",
    "S. Tang et al., “Online PCB Defect Detector on a New PCB Defect Dataset,” 2019. https://arxiv.org/abs/1902.06197",
    "S. Lv et al., “A Dataset for Deep Learning Based Detection of Printed Circuit Board Surface Defect,” Scientific Data, vol. 11, Art. no. 811, 2024. https://doi.org/10.1038/s41597-024-03656-8",
    "H. Yan et al., “Industrial Printed Circuit Board Surface Defect Dataset for Object Detection,” Scientific Data, 2026. https://doi.org/10.1038/s41597-026-07684-4",
    "J. Tang et al., “PCB-YOLO: An Improved Detection Algorithm of PCB Surface Defects Based on YOLOv5,” Sustainability, vol. 15, no. 7, Art. no. 5963, 2023. https://doi.org/10.3390/su15075963",
    "M. Yuan et al., “YOLO-HMC: An Improved Method for PCB Surface Defect Detection,” IEEE Trans. Instrum. Meas., vol. 73, pp. 1–11, Art. no. 2001611, 2024. https://doi.org/10.1109/TIM.2024.3351241",
    "Q. Li et al., “WaveCNet: Wavelet Integrated CNNs to Suppress Aliasing Effect for Noise-Robust Image Classification,” IEEE Trans. Image Process., vol. 30, pp. 7074–7089, 2021. https://doi.org/10.1109/TIP.2021.3101395",
    "Z. Qin et al., “FcaNet: Frequency Channel Attention Networks,” ICCV, 2021. https://arxiv.org/abs/2012.11879",
    "Y. Rao et al., “DynamicViT: Efficient Vision Transformers with Dynamic Token Sparsification,” NeurIPS, 2021. https://arxiv.org/abs/2106.02034",
    "T.-Y. Lin et al., “Microsoft COCO: Common Objects in Context,” ECCV, 2014. https://arxiv.org/abs/1405.0312",
]


def set_columns(section, count: int) -> None:
    sect_pr = section._sectPr
    cols = sect_pr.find(qn("w:cols"))
    if cols is None:
        cols = OxmlElement("w:cols")
        sect_pr.append(cols)
    cols.set(qn("w:num"), str(count))
    cols.set(qn("w:space"), "360")
    cols.attrib.pop(qn("w:equalWidth"), None)


def configure_section(section, columns: int) -> None:
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(0.62)
    section.right_margin = Inches(0.62)
    section.header_distance = Inches(0.5)
    section.footer_distance = Inches(0.5)
    set_columns(section, columns)


def set_run_font(run, size: float, *, bold=False, italic=False) -> None:
    run.font.name = "Times New Roman"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor(0, 0, 0)


def add_hyperlink(paragraph, text: str, url: str, *, size=9.0) -> None:
    """Add a visible black external hyperlink without changing the template look."""
    relationship = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Times New Roman")
    fonts.set(qn("w:hAnsi"), "Times New Roman")
    fonts.set(qn("w:eastAsia"), "Times New Roman")
    properties.append(fonts)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "000000")
    properties.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "none")
    properties.append(underline)
    font_size = OxmlElement("w:sz")
    font_size.set(qn("w:val"), str(round(size * 2)))
    properties.append(font_size)
    run.append(properties)
    node = OxmlElement("w:t")
    node.text = text
    run.append(node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def style_paragraph(paragraph, *, first_line=True, size=10.0) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(3)
    fmt.line_spacing = 1.15
    fmt.first_line_indent = Inches(0.2) if first_line else Inches(0)
    for run in paragraph.runs:
        set_run_font(run, size)


def body(doc, text: str, *, first=False):
    p = doc.add_paragraph()
    p.style = "Body Text"
    citation_runs = []
    citation_pattern = re.compile(r"\[\d+(?:,\s*\d+)*\](?:–\[\d+\])?")
    for part in re.split(f"({citation_pattern.pattern})", text):
        if not part:
            continue
        run = p.add_run(part)
        if citation_pattern.fullmatch(part):
            citation_runs.append(run)
    style_paragraph(p, first_line=not first)
    for run in citation_runs:
        run.font.superscript = True
    return p


def heading(doc, text: str):
    p = doc.add_paragraph()
    p.style = "Heading 1"
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text.upper())
    set_run_font(r, 10.0)
    r.font.small_caps = True
    return p


def subheading(doc, text: str):
    p = doc.add_paragraph()
    p.style = "Heading 2"
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    set_run_font(r, 10.0, italic=True)
    return p


def caption(doc, text: str, *, table=False):
    p = doc.add_paragraph()
    p.style = "table head" if table else "figure caption"
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_together = True
    p.paragraph_format.keep_with_next = table
    r = p.add_run(text)
    set_run_font(r, 8.0)
    return p


def extract_figure_blobs(doc) -> dict[int, bytes]:
    ns = {
        "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
        "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    }
    result: dict[int, bytes] = {}
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text.startswith("Figure ") or "." not in text:
            continue
        try:
            number = int(text.split()[1].rstrip("."))
        except ValueError:
            continue
        node = paragraph._p
        if not node.xpath(".//a:blip"):
            node = node.getprevious()
        while node is not None:
            blips = node.xpath(".//a:blip")
            if blips:
                rid = blips[0].get(qn("r:embed"))
                result[number] = doc.part.related_parts[rid].blob
                break
            if node.tag == qn("w:p") and "".join(node.itertext()).strip():
                break
            node = node.getprevious()
    if sorted(result) != [1, 2, 3, 4, 5, 6]:
        raise RuntimeError(f"Expected six embedded figures, found {sorted(result)}")
    return result


def extract_equations(doc):
    equations = [deepcopy(node) for node in doc._body._element.xpath(".//m:oMath")]
    if len(equations) < 4:
        raise RuntimeError(f"Expected at least four display equations, found {len(equations)}")
    return equations[:4]


def clear_body(doc) -> None:
    body_element = doc._body._element
    sect_pr = body_element.sectPr
    for child in list(body_element):
        body_element.remove(child)
    body_element.append(sect_pr)


def append_numbered_equation(doc, equation, number: int) -> None:
    """Insert editable Office math with its number outside the equation editor."""
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = (0.25, 2.65, 0.35)
    for grid_col, width in zip(table._tbl.xpath("./w:tblGrid/w:gridCol"), widths):
        grid_col.set(qn("w:w"), str(round(width * 1440)))
    tbl_pr = table._tbl.tblPr
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_pr.append(layout)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "nil")
        borders.append(element)
    tbl_pr.append(borders)
    for cell, width in zip(table.rows[0].cells, widths):
        cell.width = Inches(width)
        tc_mar = OxmlElement("w:tcMar")
        for side in ("top", "left", "bottom", "right"):
            margin = OxmlElement(f"w:{side}")
            margin.set(qn("w:w"), "0")
            margin.set(qn("w:type"), "dxa")
            tc_mar.append(margin)
        cell._tc.get_or_add_tcPr().append(tc_mar)
    math_nodes = [equation] if equation.tag == qn("m:oMath") else equation.xpath(".//m:oMath")
    if not math_nodes:
        raise RuntimeError(f"Equation {number} contains no editable Office math")
    center = table.rows[0].cells[1].paragraphs[0]
    center.style = "equation"
    center.alignment = WD_ALIGN_PARAGRAPH.CENTER
    center.paragraph_format.space_before = Pt(1)
    center.paragraph_format.space_after = Pt(1)
    center._p.append(deepcopy(math_nodes[0]))
    right = table.rows[0].cells[2].paragraphs[0]
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    right.paragraph_format.space_before = Pt(1)
    right.paragraph_format.space_after = Pt(1)
    set_run_font(right.add_run(f"({number})"), 9.5)


def add_figure(doc, blob: bytes, width: float, text: str, *, page_break=False):
    p = doc.add_paragraph()
    p.style = "figure caption"
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_together = True
    p.paragraph_format.page_break_before = page_break
    p.add_run().add_picture(BytesIO(blob), width=Inches(width))
    caption_run = p.add_run()
    caption_run.add_break()
    caption_run = p.add_run(text)
    set_run_font(caption_run, 8.0)
    return p


def set_cell_text(cell, text: str, *, bold=False, size=8.0) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.style = "table col head" if bold else "table copy"
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    set_run_font(r, size, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows, widths, caption_text):
    caption(doc, caption_text, table=True)
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    total_twips = str(round(sum(widths) * 1440))
    tbl_width = tbl_pr.find(qn("w:tblW"))
    if tbl_width is None:
        tbl_width = OxmlElement("w:tblW")
        tbl_pr.insert(0, tbl_width)
    tbl_width.set(qn("w:w"), total_twips)
    tbl_width.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    table_indent = tbl_pr.find(qn("w:tblInd"))
    if table_indent is None:
        table_indent = OxmlElement("w:tblInd")
        tbl_pr.append(table_indent)
    table_indent.set(qn("w:w"), "0")
    table_indent.set(qn("w:type"), "dxa")
    for grid_col, width in zip(table._tbl.xpath("./w:tblGrid/w:gridCol"), widths):
        grid_col.set(qn("w:w"), str(round(width * 1440)))
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "bottom", "left", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")
        borders.append(element)
    tbl_pr.append(borders)
    cell_margins = OxmlElement("w:tblCellMar")
    for edge, width in (("top", "30"), ("bottom", "30"), ("left", "40"), ("right", "40")):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:w"), width)
        element.set(qn("w:type"), "dxa")
        cell_margins.append(element)
    tbl_pr.append(cell_margins)
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, bold=True)
        table.rows[0].cells[i].width = Inches(widths[i])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
            cells[i].width = Inches(widths[i])
    group_start = None
    for row_index, row in enumerate(rows, 1):
        if row[0]:
            if group_start is not None and row_index - 1 > group_start:
                merged = table.cell(group_start, 0).merge(table.cell(row_index - 1, 0))
                merged.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            group_start = row_index
    if group_start is not None and len(rows) > group_start:
        merged = table.cell(group_start, 0).merge(table.cell(len(rows), 0))
        merged.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for row in table.rows:
        row.height = Pt(13)
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
    for row in table.rows[:-1]:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.keep_with_next = True
    return table


def format_styles(doc) -> None:
    required = {
        "paper title",
        "Author",
        "Affiliation",
        "Abstract",
        "Keywords",
        "Body Text",
        "Heading 1",
        "Heading 2",
        "figure caption",
        "table head",
        "table col head",
        "table copy",
        "references",
        "equation",
    }
    available = {style.name for style in doc.styles}
    for style_name in required - available:
        doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)

    def configure(
        name,
        size,
        *,
        bold=None,
        italic=None,
        alignment=None,
        before=0,
        after=0,
        line=1.0,
        first=None,
        keep_next=None,
        keep_together=None,
    ):
        style = doc.styles[name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.italic = italic
        style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Times New Roman")
        fmt = style.paragraph_format
        fmt.alignment = alignment
        fmt.space_before = Pt(before)
        fmt.space_after = Pt(after)
        fmt.line_spacing = line
        fmt.first_line_indent = Inches(first) if first is not None else None
        fmt.keep_with_next = keep_next
        fmt.keep_together = keep_together

    configure("Normal", 10, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    configure("paper title", 24, alignment=WD_ALIGN_PARAGRAPH.CENTER, after=6)
    configure("Author", 11, alignment=WD_ALIGN_PARAGRAPH.CENTER, before=18, after=2)
    configure("Affiliation", 9, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    configure("Abstract", 9, bold=True, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, after=10, first=0.189)
    configure("Keywords", 9, bold=True, italic=True, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, after=6, first=0.19)
    configure("Body Text", 10, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, after=3, line=1.15, first=0.2)
    configure("Heading 1", 10, alignment=WD_ALIGN_PARAGRAPH.CENTER, before=8, after=4, first=0, keep_next=True, keep_together=True)
    configure("Heading 2", 10, italic=True, alignment=WD_ALIGN_PARAGRAPH.LEFT, before=6, after=3, keep_next=True, keep_together=True)
    configure("figure caption", 8, alignment=WD_ALIGN_PARAGRAPH.CENTER, before=4, after=10, first=0)
    configure("table head", 8, alignment=WD_ALIGN_PARAGRAPH.CENTER, before=12, after=6, line=0.9)
    configure("table col head", 8, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    configure("table copy", 8, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    configure("references", 8, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, after=2.5, line=Pt(9))
    configure("equation", 10, alignment=WD_ALIGN_PARAGRAPH.CENTER, before=12, after=12, line=0.9)


def build() -> None:
    source = Document(SOURCE)
    figures = extract_figure_blobs(source)
    equations = extract_equations(source)
    doc = source
    clear_body(doc)
    format_styles(doc)

    configure_section(doc.sections[0], 1)
    title = doc.add_paragraph()
    title.style = "paper title"
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(6)
    r = title.add_run(TITLE_LINE_1)
    set_run_font(r, 24.0)
    r.add_break()
    r2 = title.add_run(TITLE_LINE_2)
    set_run_font(r2, 24.0)

    authors = doc.add_paragraph()
    authors.style = "Author"
    authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
    authors.paragraph_format.space_after = Pt(2)
    r = authors.add_run(
        "Jiefeng Liangᵃ, Lihua Luoᵃᵇ*, Sijin Tanᵃ, Yanni Linᵃ, "
        "Zhizhuo Zhaoᵃ, and Zhaofeng Caiᵃ"
    )
    set_run_font(r, 11.0)
    for line in (
        "ᵃDepartment of Computing, Guangdong University of Science and Technology, Dongguan 523083, China",
        "ᵇFaculty of Computer and Mathematical Sciences, Universiti Teknologi MARA, 40450 Shah Alam, Selangor, Malaysia",
        "*Corresponding author: Lihua Luo, 260423630@qq.com",
    ):
        r.add_break()
        r = authors.add_run(line)
        set_run_font(r, 9.0)
    r.add_break()
    r = authors.add_run("Project repository: ")
    set_run_font(r, 9.0)
    add_hyperlink(authors, PROJECT_URL, PROJECT_URL, size=9.0)

    two_col = doc.add_section(WD_SECTION.CONTINUOUS)
    configure_section(two_col, 2)

    p = doc.add_paragraph()
    p.style = "Abstract"
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run("Abstract—")
    set_run_font(r, 9.0, bold=True, italic=True)
    r = p.add_run(ABSTRACT)
    set_run_font(r, 9.0, bold=True)

    p = doc.add_paragraph()
    p.style = "Keywords"
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run("Keywords—")
    set_run_font(r, 9.0, bold=True, italic=True)
    r = p.add_run("image processing, PCB defect detection, sparse routing, wavelets, false alarms")
    set_run_font(r, 9.0, bold=True)

    heading(doc, "I. Introduction")
    body(doc, "Automated optical inspection must detect small, low-contrast defects among repeated traces without alarming on normal boards. Positive-only benchmarks omit that operational risk, while route maps and FLOPs alone cannot establish accuracy, speed, or reliability.", first=True)
    body(doc, "WSR is therefore tested as an auditable probe. Fixed Haar responses rank P3 locations, and gather–refine–scatter transforms only k=round(ρHW) tokens. Paired seeds, same-budget and matched-capacity controls, device timing, and defect-free inputs test route enrichment, held-out accuracy, realized latency, and cue specificity.")
    body(doc, "Our exact-budget probe separates arithmetic savings from measured speed and exposes normal-board false alarms. Because the primary gain is nonsignificant and no variant passes the joint accuracy–latency gate, we report the evidence boundary rather than claim superiority.")

    heading(doc, "II. Related Work")
    body(doc, "Detectors span region proposals [1], set prediction [2], sparse proposals [3], and compact YOLO/DETR models [4], [5]. PCB protocols also differ: DeepPCB provides target/template pairs [6], DsPCBSD+ has nine AOI classes [7], and PCB-IND adds illumination and long-tail variation [8]. Results from PCB-YOLO and YOLO-HMC therefore remain protocol-dependent [9], [10].", first=True)
    body(doc, "Wavelet anti-aliasing and frequency attention [11], [12] motivate WSR, while dynamic token methods allocate input-dependent computation [13]. Yet scoring, indexing, and memory movement may dominate skipped arithmetic, requiring direct timing.")

    heading(doc, "III. Wavelet-Conditioned Sparse Routing")
    body(doc, "Let X∈R^(C×H×W) be the P3/8 feature, split into wave and refinement tensors Xw and Xr. Fixed grouped Haar convolution gives", first=True)
    append_numbered_equation(doc, equations[0], 1)
    body(doc, "Normalized detail energy Ehf, low-frequency magnitude Ell, and local residual Gll enter a small router:", first=True)
    append_numbered_equation(doc, equations[1], 2)
    body(doc, "Flattening spatial locations, the router selects an exact budget", first=True)
    append_numbered_equation(doc, equations[2], 3)
    body(doc, "A normalized 3×3 kernel contextualizes Xr; only selected indices are gathered. For token tensor T,", first=True)
    append_numbered_equation(doc, equations[3], 4)
    body(doc, "Unrouted tokens remain unchanged. A softmax gate fuses both branches with an identity skip. One block is inserted at P3 with ρ=0.25; confidence factors and detector losses train the router although hard indices receive no gradient (Figure 1).", first=True)

    heading(doc, "IV. Experiments")
    subheading(doc, "A. Protocol")
    body(doc, "Table 1 gives the frozen splits. DsPCBSD+ uses seven paired seeds; DeepPCB uses three descriptive pairs and audits 500 clean templates separately. Models share the YOLO11s checkpoint, 640-pixel input, 80-epoch limit, deterministic SGD, and augmentation. A pinned COCO evaluator [14] reports AP; RTX 3090 timing uses synchronized ABBA/BAAB cycles.", first=True)

    wide = doc.add_section(WD_SECTION.CONTINUOUS)
    configure_section(wide, 1)
    add_figure(doc, figures[1], 7.05, "Figure 1. Compact WSR flow. Fixed Haar cues drive wave gating and dense routing; only the gathered channel MLP uses the exact ρHW budget.")
    two_col = doc.add_section(WD_SECTION.CONTINUOUS)
    configure_section(two_col, 2)

    body(doc, "Pretraining transfer must exceed 99%; paths, labels, boxes, hashes, and similarity are audited before training. P3/25% is selected on validation data before locked-test access; later studies are explicitly diagnostic or exploratory.")
    add_table(
        doc,
        ("Dataset", "Split", "Images", "Boxes"),
        (
            ("DsPCBSD+", "Train", "7,387", "14,590"),
            ("", "Validation", "821", "1,594"),
            ("", "Locked test", "2,051", "4,092"),
            ("DeepPCB", "Train", "850", "5,870"),
            ("", "Validation", "150", "1,003"),
            ("", "Official test", "500", "3,140"),
        ),
        (0.88, 0.85, 0.62, 0.62),
        "Table 1. Frozen dataset partitions after conversion and validation.",
    )
    body(doc, "Route recall is |Mi∩Bi|/|Bi|; enrichment divides it by the selected fraction |Mi|/Ni. Board-FPR is the share of clean templates with a retained detection.")

    subheading(doc, "B. Accuracy")
    body(doc, "Table 2 reports architecture-controlled tests. DsPCBSD+ gains only 0.155 AP across four of seven wins (p=0.8125; 95% interval [−0.535, 0.844]). DeepPCB gains 4.196 points, but n=3 and p=0.25 preclude confirmation.", first=True)
    add_table(
        doc,
        ("Dataset", "Model", "n", "AP50:95", "AP50"),
        (
            ("DsPCBSD+", "YOLO11s", "7", "46.54±0.36", "80.25±0.48"),
            ("", "WSR-YOLO11s", "7", "46.69±0.72", "80.46±0.86"),
            ("DeepPCB", "YOLO11s", "3", "64.61±1.91", "94.34±1.07"),
            ("", "WSR-YOLO11s", "3", "68.81±4.12", "95.46±0.43"),
        ),
        (0.62, 0.9, 0.25, 0.78, 0.68),
        "Table 2. Architecture-controlled test AP (%), mean±sample SD.",
    )
    body(doc, "Figure 2 shows unstable DsPCBSD+ seed effects (−0.99 to +0.98; two classes decline), whereas all DeepPCB classes gain but reuse only three checkpoints. Figure 3 gives model-blind examples. Larger published detectors gain 2.75–6.39 AP with 1.75–3.55× the parameters, under different recipes.")
    body(doc, "Scale-specific changes are also inconsistent: the effect is neither seed-stable on the primary dataset nor uniform across object sizes. Class and scale slices remain diagnostic decompositions, not independent samples.")
    wide = doc.add_section(WD_SECTION.CONTINUOUS)
    configure_section(wide, 1)
    add_figure(doc, ACCURACY_FIGURE.read_bytes(), 7.05, "Figure 2. Paired and classwise accuracy differences. Dots are paired seeds; bars show per-class mean changes, with orange indicating regressions.")
    two_col = doc.add_section(WD_SECTION.CONTINUOUS)
    configure_section(two_col, 2)
    add_figure(doc, QUALITATIVE_FIGURE.read_bytes(), 3.45, "Figure 3. Model-blind DsPCBSD+ examples. Green: ground truth; blue/orange: true positives; dashed red: false positives. TP/FP/FN use class-matched IoU ≥0.5.")

    subheading(doc, "C. Routing, Controls, and Cost")
    body(doc, "At 25%, routes recall 79.11%/88.68% of box cells on DsPCBSD+/DeepPCB (3.164×/3.547× enrichment). Figure 4 shows that spatial priors explain much of the primary enrichment; cross-image shuffling confirms image dependence but not wavelet specificity. Equal fusion, no-HF, and matched convolution reach 48.314, 48.296, and 48.057 AP versus 48.025 for WSR, leaving capacity and optimization as plausible explanations.", first=True)
    body(doc, "The validation gate admitted WSR at +1.283 AP and 2.469× enrichment, but locked timing reversed the preliminary speed advantage.")
    body(doc, "Across 27 validation jobs, WSR reaches 48.025 AP versus 46.995 for the repeated baseline, with a 95% interval of [−0.158, 2.219]. The matched controls therefore do not isolate a unique wavelet mechanism.")
    add_figure(doc, CONTEXT_FIGURE.read_bytes(), 3.45, "Figure 4. Model-size and route context. (a) DsPCBSD+ AP versus parameters. (b) Same-budget route enrichment; the dashed line is uniform expectation.")
    body(doc, "WSR adds 13,342 parameters (0.14%) and 0.038 GFLOPs (0.18%), yet latency is 1.184–1.206× baseline; every paired cycle interval excludes one. Dense support and unfused indexing therefore yield sparse arithmetic, not acceleration (Table 3).")
    body(doc, "The 95% ratio intervals are [1.166, 1.203]/[1.162, 1.227] on DsPCBSD+ FP32/FP16 and [1.193, 1.220]/[1.166, 1.204] on DeepPCB, confirming slowdown in all four settings.")
    add_table(
        doc,
        ("Study", "Variant/setting", "ΔAP", "Ratio"),
        (
            ("Control", "WSR P3/25%", "+1.030", "—"),
            ("", "no HF cue", "+1.301", "—"),
            ("", "equal fusion", "+1.319", "—"),
            ("", "matched convolution", "+1.062", "—"),
            ("Latency", "DsPCBSD+ FP32/FP16", "—", "1.184/1.194"),
            ("", "DeepPCB FP32/FP16", "—", "1.206/1.185"),
        ),
        (0.52, 1.45, 0.52, 0.72),
        "Table 3. Compact validation and latency audit. Ratios are WSR/baseline.",
    )

    subheading(doc, "D. Operational False Alarms")
    body(doc, "DeepPCB’s 500 clean templates expose risk hidden by positive-only AP. At confidence 0.25, seed-13 WSR alarms on 56.6% of boards versus 48.0% for YOLO11s (Table 4). Negative-aware training lowers the three-seed board-FPR from 53.5% to 27.5% at 0.946 recall and raises AP from 68.81 to 72.44. The exploratory paired-reference policy reaches 1.4% board-FPR, AP 72.63, and 0.959 F1, but is post-hoc, single-seed, and changes the input contract (Figure 5).", first=True)
    add_table(
        doc,
        ("Model", "0.05", "0.10", "0.25", "0.50"),
        (
            ("YOLO11s", "79.0", "66.2", "48.0", "28.0"),
            ("WSR-YOLO11s", "87.6", "77.0", "56.6", "35.0"),
        ),
        (1.15, 0.48, 0.48, 0.48, 0.48),
        "Table 4. Seed-13 board-FPR (%) on 500 clean DeepPCB templates.",
    )
    body(doc, "Target-only mitigation preserves the sensor input; paired assistance requires a registered template. Validation-calibrated 1% caps transfer to 1.8%–6.2% test board-FPR with 0.790 mean recall, so production negatives remain necessary.")
    add_figure(doc, FALSE_ALARM_FIGURE.read_bytes(), 3.45, "Figure 5. DeepPCB negative-template audit. (a) Target-only confidence sweep. (b) Recall–board-FPR trade-off; paired input changes the inference contract.")

    subheading(doc, "E. Auxiliary Robustness Evidence")
    body(doc, "Across eight corruptions and five severities, WSR averages 46.51 AP versus 42.68 and wins 24 of 40 conditions. Gains cluster in brightness, contrast, Gaussian noise, and JPEG, while three corruptions regress; clean-normalized retention is 65.6% versus 64.1%.", first=True)
    body(doc, "Low-/high-only reconstruction gives 67.43/0.95 AP for WSR and 66.22/1.67 for baseline. Removing LH, HL, or HH costs WSR 1.32, 0.70, or 1.68 points versus 1.59, 0.92, or 3.68 for baseline. Because the interventions modify detector input, they are mechanistic evidence rather than proof of internal Haar causality (Figure 6).")
    add_figure(doc, ROBUSTNESS_FIGURE.read_bytes(), 3.45, "Figure 6. DeepPCB seed-13 diagnostics. (a) Mean WSR-minus-baseline AP under corruptions. (b) AP after matched input-Haar interventions.")

    heading(doc, "V. Findings and Limitations")
    body(doc, "Route enrichment is not an accuracy surrogate, arithmetic sparsity is not systems efficiency, and positive-only AP hides normal-board alarms. A valid audit combines paired held-out accuracy, device timing, negative operating points, same-budget controls, matched capacity, and source-bound records. Seeds and latency cycles remain the sampling units; class, corruption, and threshold slices are decompositions rather than new samples.", first=True)

    subheading(doc, "A. Image-Processing Role and Deployment Gate")
    body(doc, "Filtering, binarization, and edge operators remain useful for normalization, registration, and ROI masks, but fixed thresholds are brittle under illumination and repeated copper patterns. WSR-YOLO instead learns multiscale context. Deployment freezes classwise thresholds under a validation board-FPR ceiling and reports locked-test AP, recall, board-FPR uncertainty, and end-to-end latency; camera or layout changes trigger recalibration, while loss of registration disables the paired branch.", first=True)
    body(doc, "Limits include three DeepPCB pairs, one GPU, incomplete comparator timing, hard top-k, and absent board/lot IDs. Reference assistance requires preregistration and independent production negatives. Runs bind Python 3.10, PyTorch 2.5.1/CUDA 12.1, Ultralytics 8.4.50, source revision, manifest, seed, checkpoint, and evaluator output; validation controls never read test labels.")

    heading(doc, "VI. Conclusion")
    body(doc, "WSR concentrates selected P3 cells inside defects at little arithmetic cost, yet establishes neither a reliable primary accuracy gain nor acceleration, and target-only false alarms remain high. Its contribution is an auditable image-processing result and evaluation template. Future work should test fused sparse kernels, more paired seeds, group-safe splits, and independent reference-assisted evaluation.", first=True)

    heading(doc, "ACKNOWLEDGEMENTS")
    body(doc, ACKNOWLEDGEMENT, first=True)

    heading(doc, "REFERENCES")
    for number, reference in enumerate(REFERENCES, 1):
        p = doc.add_paragraph()
        p.style = "references"
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.left_indent = Inches(0.18)
        p.paragraph_format.first_line_indent = Inches(-0.18)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2.5)
        p.paragraph_format.line_spacing = Pt(9)
        metadata, url = reference.rsplit(" ", 1)
        r = p.add_run(f"[{number}]  {metadata} ")
        set_run_font(r, 8.0)
        add_hyperlink(p, url, url, size=8.0)

    doc.core_properties.title = f"{TITLE_LINE_1} {TITLE_LINE_2}"
    doc.core_properties.subject = "Compact four-page conference manuscript"
    doc.save(OUTPUT)
    print(f"Built: {OUTPUT}")
    print(f"Main:  {SOURCE}")


if __name__ == "__main__":
    build()
